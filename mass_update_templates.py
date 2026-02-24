#!/usr/bin/env python3
"""
Mass find/replace in Word .docx templates using an Excel lookup table.

Notes:
- Uses python-docx (works with .docx, not legacy .doc).
- Replacement is run-by-run to preserve formatting. Matches that span
  multiple runs may not be replaced.
"""

from __future__ import annotations

import argparse
import csv
import re
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Tuple

from docx import Document
from openpyxl import load_workbook


DEFAULT_EXCEL_PATH = (
    r"C:\Users\Tim\OneDrive - quillarrowlaw.com\Documents\ClioTemplates_CustomFields_MassUpdate"
    r"\CustomField_LookupTable.xlsx"
)
DEFAULT_SHEET_NAME = "LookupTable"

@dataclass(frozen=True)
class Replacement:
    old_text: str
    new_text: str
    pattern: re.Pattern[str]

# Normalize the header value to a lowercase string
def _normalize_header(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip().lower()


# Resolve the column index based on the header value
def _resolve_column_index(
    header_row: Iterable[object],
    header_map: dict[str, int],
    column_arg: str | None,
    fallbacks: List[str],
    column_label: str,
) -> int:
    if column_arg:
        if column_arg.strip().isdigit():
            idx = int(column_arg.strip()) - 1
            if idx < 0 or idx >= len(header_row):
                raise ValueError(
                    f"{column_label} column index {column_arg} is out of range."
                )
            return idx
        normalized = _normalize_header(column_arg)
        if normalized in header_map:
            return header_map[normalized]
        raise ValueError(
            f"{column_label} column '{column_arg}' was not found in the header row."
        )

    for name in fallbacks:
        if name in header_map:
            return header_map[name]

    # Fallback: first two columns
    if len(header_row) >= 2:
        return 0 if column_label == "Old" else 1

    raise ValueError(
        "Lookup table must contain at least two columns for Old/New values."
    )


# Load the lookup table from the Excel file
# The lookup table is a Excel file with two columns: Old and New
# The Old column is the value to be replaced
# The New column is the value to replace the Old value with
# The lookup table is used to replace the Old value with the New value in the Word .docx templates
def load_lookup_table(  
    excel_path: Path,
    sheet_name: str,
    old_col: str | None,
    new_col: str | None,
    use_regex: bool,
    ignore_case: bool,
) -> List[Replacement]:
    if not excel_path.exists():
        raise FileNotFoundError(f"Excel file not found: {excel_path}")

    # Load the workbook from the Excel file
    workbook = load_workbook(excel_path, data_only=True, read_only=True)

    # Check if the sheet name is valid
    if sheet_name not in workbook.sheetnames:
        raise ValueError(
            f"Worksheet '{sheet_name}' not found. Available: {workbook.sheetnames}"
        )

    # Get the sheet from the workbook
    sheet = workbook[sheet_name]

    # Iterate through the rows of the sheet
    rows = sheet.iter_rows(values_only=True)
    try:
        # Get the header row from the rows
        header_row = next(rows)
    
    # If the header row is not found, raise an error
    except StopIteration as exc:
        raise ValueError("Lookup table sheet is empty.") from exc

    # Create a map of the header values to their column indices
    header_map: dict[str, int] = {}
    for idx, value in enumerate(header_row):
        normalized = _normalize_header(value)
        if normalized:
            header_map[normalized] = idx

    # Resolve the column index for the Old value
    old_idx = _resolve_column_index(
        header_row,
        header_map,
        old_col,
        ["old", "old value", "old_value", "old values", "oldvalues"],
        "Old",
    )

    # Resolve the column index for the New value
    new_idx = _resolve_column_index(
        header_row,
        header_map,
        new_col,
        ["new", "new value", "new_value", "new values", "newvalues"],
        "New",
    )

    # Create a list of replacement tuples
    # The flags are used to determine if the replacement should be case-insensitive
    flags = re.IGNORECASE if ignore_case else 0

    # Create a list of replacement entries
    replacements: List[Replacement] = []

    # Iterate through the rows of the sheet
    for row in rows:
        # If the row is None, continue
        if row is None:
            continue
        # Get the Old value from the row
        old_value = row[old_idx] if old_idx < len(row) else None
        # Get the New value from the row
        new_value = row[new_idx] if new_idx < len(row) else None
        # If the Old value is None or empty, continue
        if old_value is None or str(old_value).strip() == "":
            continue

        # Convert the Old value to a string
        old_text = str(old_value)
        # Convert the New value to a string
        new_text = "" if new_value is None else str(new_value)
        # Convert the Old value to a regex pattern
        pattern_text = old_text if use_regex else re.escape(old_text)
        # Create a regex pattern from the Old value

        try:
            # Create a regex pattern from the Old value
            pattern = re.compile(pattern_text, flags=flags)
        # If the regex pattern is invalid, raise an error
        except re.error as exc:
            # If the regex pattern is invalid, raise an error
            raise ValueError(
                f"Invalid regex pattern in Old value '{old_text}': {exc}"
            ) from exc

        # Add the replacement entry to the list of replacements
        replacements.append(
            Replacement(old_text=old_text, new_text=new_text, pattern=pattern)
        )

    # If there are no replacements, raise an error
    if not replacements:
        # If there are no replacements, raise an error
        raise ValueError("No replacement rows found in the lookup table.")

    # Return the list of replacements
    return replacements


# Iterate through the paragraphs in the container
def iter_paragraphs(container) -> Iterable:
    # Iterate through the paragraphs in the container
    for paragraph in container.paragraphs:
        yield paragraph
    # Iterate through the tables in the container
    for table in container.tables:
        # Iterate through the rows in the table
        for row in table.rows:
            # Iterate through the cells in the row
            for cell in row.cells:
                yield from iter_paragraphs(cell)


# Apply the replacements to the text
def apply_replacements(
    text: str, replacements: List[Replacement], counts: List[int]
) -> Tuple[str, int]:
    # Initialize the total count of replacements
    total = 0
    # Initialize the updated text
    updated = text
    for idx, replacement in enumerate(replacements):
        # Apply the replacements to the text
        updated, count = replacement.pattern.subn(replacement.new_text, updated)
        # Track per-replacement counts
        if count:
            counts[idx] += count
            total += count
    # Return the updated text and the total count of replacements
    return updated, total


# Find the run index for a character position
def _find_run_index(run_spans: List[Tuple[int, int]], position: int) -> int:
    for idx, (start, end) in enumerate(run_spans):
        if start <= position < end:
            return idx
    for idx in range(len(run_spans) - 1, -1, -1):
        if run_spans[idx][0] != run_spans[idx][1]:
            return idx
    return 0


# Replace the text in the paragraph
def replace_in_paragraph(
    paragraph, replacements: List[Replacement], counts: List[int]
) -> int:
    # Initialize the total count of replacements
    total = 0
    # Iterate through the runs in the paragraph
    for run in paragraph.runs:
        # If the run text is None, continue
        if not run.text:
            continue
        # Apply the replacements to the text
        updated, count = apply_replacements(run.text, replacements, counts)
        # If the count of replacements is greater than 0, update the run text
        if count:
            run.text = updated
            # Increment the total count of replacements
            total += count
    # Return the total count of replacements
    return total


# Replace the text across runs in the paragraph
def replace_in_paragraph_join_runs(
    paragraph, replacements: List[Replacement], counts: List[int]
) -> int:
    total = 0
    safety_limit = 10000
    iterations = 0

    while iterations < safety_limit:
        runs = paragraph.runs
        run_texts = [run.text or "" for run in runs]
        if not any(run_texts):
            break
        full_text = "".join(run_texts)

        earliest_match = None
        earliest_index = -1
        for idx, replacement in enumerate(replacements):
            match = replacement.pattern.search(full_text)
            if match and (
                earliest_match is None or match.start() < earliest_match.start()
            ):
                earliest_match = match
                earliest_index = idx

        if earliest_match is None:
            break

        start, end = earliest_match.span()
        if start == end:
            break

        run_spans: List[Tuple[int, int]] = []
        pos = 0
        for text in run_texts:
            run_spans.append((pos, pos + len(text)))
            pos += len(text)

        start_run = _find_run_index(run_spans, start)
        end_run = _find_run_index(run_spans, end - 1)

        start_run_start = run_spans[start_run][0]
        end_run_start = run_spans[end_run][0]

        prefix = run_texts[start_run][: start - start_run_start]
        suffix = run_texts[end_run][end - end_run_start :]

        runs[start_run].text = (
            prefix + replacements[earliest_index].new_text + suffix
        )
        for idx in range(start_run + 1, end_run + 1):
            runs[idx].text = ""

        counts[earliest_index] += 1
        total += 1
        iterations += 1

    return total


# Process the document
def process_document(
    input_path: Path,
    output_path: Path,
    replacements: List[Replacement],
    include_headers_footers: bool,
    dry_run: bool,
    join_runs: bool,
) -> Tuple[bool, int, List[int]]:
    # Load the document from the input path
    doc = Document(str(input_path))
    # Initialize the total count of replacements
    total_replacements = 0
    # Initialize the per-replacement counts
    counts = [0] * len(replacements)
    # Define a function to process the container
    def process_container(container) -> None:
        nonlocal total_replacements
        # Iterate through the paragraphs in the container
        for paragraph in iter_paragraphs(container):
            # Replace the text in the paragraph
            if join_runs:
                total_replacements += replace_in_paragraph_join_runs(
                    paragraph, replacements, counts
                )
            else:
                total_replacements += replace_in_paragraph(
                    paragraph, replacements, counts
                )

    # Process the document
    process_container(doc)

    # If include headers and footers, process the headers and footers
    if include_headers_footers:
        # Initialize a set of seen elements
        seen = set()
        for section in doc.sections:
            # Get the header and footer groups for the section
            header_footer_groups = [
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ]
            # Iterate through the header and footer groups
            for header_footer in header_footer_groups:
                # Get the element id
                element_id = id(header_footer._element)
                # If the element id is in the set of seen elements, continue
                if element_id in seen:
                    # If the element id is in the set of seen elements, continue
                    continue
                # Add the element id to the set of seen elements
                seen.add(element_id)
                # Process the container
                process_container(header_footer)

    # If the total count of replacements is greater than 0, set the changed flag to True
    changed = total_replacements > 0
    # If dry run is enabled, return the changed flag and the total count of replacements
    if dry_run:
        return changed, total_replacements, counts

    # Create the output directory if it doesn't exist
    output_path.parent.mkdir(parents=True, exist_ok=True)
    # If the changed flag is True, save the document to the output path
    if changed:
        doc.save(str(output_path))
    else:
        # If the changed flag is False, copy the input path to the output path
        if output_path.resolve() != input_path.resolve():
            # Copy the input path to the output path
            shutil.copy2(input_path, output_path)
    # Return the changed flag and the total count of replacements

    return changed, total_replacements, counts


def main() -> int:
    # Create the parser
    parser = argparse.ArgumentParser(
        description="Mass find/replace in Word .docx templates using an Excel lookup table."
    )
    # Add the excel argument
    parser.add_argument(
        "--excel",
        default=DEFAULT_EXCEL_PATH,
        help="Path to Excel lookup table (.xlsx).",
    )
    # Add the sheet argument
    parser.add_argument(
        "--sheet",
        default=DEFAULT_SHEET_NAME,
        help="Worksheet name containing Old/New values.",
    )
    # Add the input-dir argument
    parser.add_argument(
        "--input-dir",
        required=True,
        help="Directory containing Word .docx templates.",
    )
    # Add the output-dir argument
    parser.add_argument(
        "--output-dir",
        help=(
            "Directory to write updated templates. "
            "Defaults to <input-dir>_updated if not using --in-place."
        ),
    )
    # Add the in-place argument
    parser.add_argument(
        "--in-place",
        action="store_true",
        help="Overwrite templates in place (use with caution).",
    )
    # Add the old-col argument
    parser.add_argument(
        "--old-col",
        help="Old value column header or 1-based index (e.g., 'Old' or '1').",
    )
    # Add the new-col argument
    parser.add_argument(
        "--new-col",
        help="New value column header or 1-based index (e.g., 'New' or '2').",
    )
    # Add the literal argument
    parser.add_argument(
        "--literal",
        action="store_true",
        help="Treat Old values as literal text instead of regex patterns.",
    )
    # Add the ignore-case argument
    parser.add_argument(
        "--ignore-case",
        action="store_true",
        help="Case-insensitive matching.",
    )
    # Add the skip-headers-footers argument
    parser.add_argument(
        "--skip-headers-footers",
        action="store_true",
        help="Skip replacements in headers and footers.",
    )
    # Add the dry-run argument
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Preview matches without writing files.",
    )
    parser.add_argument(
        "--join-runs",
        action="store_true",
        help=(
            "Allow replacements spanning multiple Word runs. "
            "Uses the first run's formatting for the replaced span."
        ),
    )
    parser.add_argument(
        "--report",
        help=(
            "CSV report path for per-document replacements. "
            "Defaults to replacement_report.csv in the output directory."
        ),
    )
    # Parse the arguments
    args = parser.parse_args()
    # Resolve the input directory

    input_dir = Path(args.input_dir).resolve()
    # If the input directory does not exist or is not a directory, raise an error
    if not input_dir.exists() or not input_dir.is_dir():
        # If the input directory does not exist or is not a directory, raise an error
        print(f"Input directory not found: {input_dir}", file=sys.stderr)
        return 1

    # Initialize the output directory
    output_dir: Path
    # If in-place is enabled, set the output directory to the input directory
    if args.in_place:
        output_dir = input_dir
    else:
        # If output-dir is not set, set the output directory to the input directory with _updated appended
        output_dir = (
            Path(args.output_dir).resolve()
            if args.output_dir
            else Path(f"{input_dir}_updated")
        )

    # Set the report path
    report_path = (
        Path(args.report).resolve()
        if args.report
        else (input_dir if args.dry_run else output_dir) / "replacement_report.csv"
    )

    # Load the lookup table
    replacements = load_lookup_table(
        # Resolve the excel path
        Path(args.excel).resolve(),
        # Get the sheet name
        args.sheet,
        # Get the old column
        args.old_col,
        args.new_col,
        # Use the regex flag
        use_regex=not args.literal,
        ignore_case=args.ignore_case,
    )

    # Get the list of docx files
    docx_files = [
        path
        # Iterate through the files in the input directory
        for path in input_dir.rglob("*.docx")
        # If the file name starts with ~$, continue
        if not path.name.startswith("~$")
    ]

    # If there are no docx files, raise an error
    if not docx_files:
        # If there are no docx files, raise an error
        print(f"No .docx files found under {input_dir}", file=sys.stderr)
        return 1

    # Print the number of replacement rows
    print(f"Loaded {len(replacements)} replacement rows.")
    # Print the number of docx files
    print(f"Processing {len(docx_files)} .docx files from {input_dir}")
    # If dry run is enabled, print a message
    if args.dry_run:
        # If dry run is enabled, print a message
        print("Dry run enabled; no files will be written.")

    # Initialize the total count of files changed
    total_files_changed = 0
    # Initialize the total count of replacements
    total_replacements = 0
    # Initialize the report rows
    report_rows: List[Tuple[str, str, str, int]] = []

    for file_path in docx_files:
        # Get the relative path
        rel_path = file_path.relative_to(input_dir)
        # Get the target path
        target_path = output_dir / rel_path
        # Process the document
        changed, count, counts = process_document(
            # Get the file path
            file_path,
            target_path,
            # Get the replacements
            replacements,
            # If skip headers and footers is enabled, include headers and footers
            include_headers_footers=not args.skip_headers_footers,
            # If dry run is enabled, dry run the document
            dry_run=args.dry_run,
            join_runs=args.join_runs,
        )
        # Increment the total count of replacements
        total_replacements += count
        # Add per-replacement counts to the report
        for idx, rep_count in enumerate(counts):
            if rep_count:
                replacement = replacements[idx]
                report_rows.append(
                    (
                        str(rel_path),
                        replacement.old_text,
                        replacement.new_text,
                        rep_count,
                    )
                )
        # If the document was changed, increment the total count of files changed
        if changed:
            # Increment the total count of files changed
            total_files_changed += 1
            # Set the status to updated
            status = "updated"
        # If the document was not changed, set the status to no changes
        else:
            # Set the status to no changes
            status = "no changes"
        # Print the relative path, status, and number of replacements
        print(f"{rel_path} -> {status} ({count} replacements)")

    # Print the total count of files changed, total count of replacements, and output directory
    print(
        f"Done. Files updated: {total_files_changed}/{len(docx_files)}. "
        f"Total replacements: {total_replacements}."
    )
    # If in-place is not enabled and dry run is not enabled, print the output directory
    if not args.in_place and not args.dry_run:
        # Print the output directory
        print(f"Output directory: {output_dir}")

    # Write the report to the output path
    report_path.parent.mkdir(parents=True, exist_ok=True)
    with report_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["document", "old_value", "new_value", "count"])
        writer.writerows(report_rows)
    print(f"Report written to: {report_path}")

    return 0

# If the script is run directly, raise an error
if __name__ == "__main__":
    # Raise an error
    raise SystemExit(main())
