#!/usr/bin/env python3
"""
Mass find/replace in Word .docx templates using an Excel lookup table.

Notes:
- Uses python-docx (works with .docx, not legacy .doc).
- Replacement is run-by-run to preserve formatting. Matches that span
  multiple runs may not be replaced.
"""

from __future__ import annotations

import argparse
import csv
import re
import shutil
import sys
import time
import uuid
import zipfile
from html import escape, unescape
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Sequence, Tuple

from docx import Document
from openpyxl import load_workbook


DEFAULT_EXCEL_PATH = (
    r"C:\Users\Tim\OneDrive - quillarrowlaw.com\Documents\ClioTemplates_CustomFields_MassUpdate"
    r"\CustomField_LookupTable.xlsx"
)
DEFAULT_SHEET_NAME = "LookupTable"

@dataclass(frozen=True)
class Replacement:
    old_text: str
    new_text: str
    pattern: re.Pattern[str]

# Normalize the header value to a lowercase string
def _normalize_header(value: object) -> str:
    return "" if value is None else str(value).strip().lower()


# Resolve the column index based on the header value
def _resolve_column_index(
    header_row: Sequence[object],
    header_map: dict[str, int],
    column_arg: str | None,
    fallbacks: List[str],
    column_label: str,
) -> int:
    if column_arg:
        if column_arg.strip().isdigit():
            idx = int(column_arg.strip()) - 1
            if idx < 0 or idx >= len(header_row):
                raise ValueError(
                    f"{column_label} column index {column_arg} is out of range."
                )
            return idx
        normalized = _normalize_header(column_arg)
        if normalized in header_map:
            return header_map[normalized]
        raise ValueError(
            f"{column_label} column '{column_arg}' was not found in the header row."
        )

    for name in fallbacks:
        if name in header_map:
            return header_map[name]

    # Fallback: first two columns
    if len(header_row) >= 2:
        return 0 if column_label == "Old" else 1

    raise ValueError(
        "Lookup table must contain at least two columns for Old/New values."
    )


# Load the lookup table from the Excel file
# The lookup table is a Excel file with two columns: Old and New
# The Old column is the value to be replaced
# The New column is the value to replace the Old value with
# The lookup table is used to replace the Old value with the New value in the Word .docx templates
def load_lookup_table(  
    excel_path: Path,
    sheet_name: str,
    old_col: str | None,
    new_col: str | None,
    use_regex: bool,
    ignore_case: bool,
) -> List[Replacement]:
    if not excel_path.exists():
        raise FileNotFoundError(f"Excel file not found: {excel_path}")

    # Load the workbook from the Excel file
    workbook = load_workbook(excel_path, data_only=True, read_only=True)

    # Check if the sheet name is valid
    if sheet_name not in workbook.sheetnames:
        raise ValueError(
            f"Worksheet '{sheet_name}' not found. Available: {workbook.sheetnames}"
        )

    # Get the sheet from the workbook
    sheet = workbook[sheet_name]

    # Iterate through the rows of the sheet
    rows = sheet.iter_rows(values_only=True)
    try:
        # Get the header row from the rows
        header_row = next(rows)

    # If the header row is not found, raise an error
    except StopIteration as exc:
        raise ValueError("Lookup table sheet is empty.") from exc

    # Create a map of the header values to their column indices
    header_map: dict[str, int] = {}
    for idx, value in enumerate(header_row):
        if normalized := _normalize_header(value):
            header_map[normalized] = idx

    # Resolve the column index for the Old value
    old_idx = _resolve_column_index(
        header_row,
        header_map,
        old_col,
        ["old", "old value", "old_value", "old values", "oldvalues"],
        "Old",
    )

    # Resolve the column index for the New value
    new_idx = _resolve_column_index(
        header_row,
        header_map,
        new_col,
        ["new", "new value", "new_value", "new values", "newvalues"],
        "New",
    )

    # Create a list of replacement tuples
    # The flags are used to determine if the replacement should be case-insensitive
    flags = re.IGNORECASE if ignore_case else 0

    # Create a list of replacement entries
    replacements: List[Replacement] = []

    # Iterate through the rows of the sheet
    for row in rows:
        # If the row is None, continue
        if row is None:
            continue
        # Get the Old value from the row
        old_value = row[old_idx] if old_idx < len(row) else None
        # Get the New value from the row
        new_value = row[new_idx] if new_idx < len(row) else None
        # If the Old value is None or empty, continue
        if old_value is None or not str(old_value).strip():
            continue

        # Convert the Old value to a string
        old_text = str(old_value)
        # Convert the New value to a string
        new_text = "" if new_value is None else str(new_value)
        # Convert the Old value to a regex pattern
        pattern_text = old_text if use_regex else re.escape(old_text)
        # Create a regex pattern from the Old value

        try:
            # Create a regex pattern from the Old value
            pattern = re.compile(pattern_text, flags=flags)
        # If the regex pattern is invalid, raise an error
        except re.error as exc:
            # If the regex pattern is invalid, raise an error
            raise ValueError(
                f"Invalid regex pattern in Old value '{old_text}': {exc}"
            ) from exc

        # Add the replacement entry to the list of replacements
        replacements.append(
            Replacement(old_text=old_text, new_text=new_text, pattern=pattern)
        )

    # If there are no replacements, raise an error
    if not replacements:
        # If there are no replacements, raise an error
        raise ValueError("No replacement rows found in the lookup table.")

    # Return the list of replacements
    return replacements


# Iterate through the paragraphs in the container
def iter_paragraphs(container) -> Iterable:
    # Iterate through the paragraphs in the container
    yield from container.paragraphs
    # Iterate through the tables in the container
    for table in container.tables:
        # Iterate through the rows in the table
        for row in table.rows:
            # Iterate through the cells in the row
            for cell in row.cells:
                yield from iter_paragraphs(cell)


# Apply the replacements to the text
def apply_replacements(
    text: str, replacements: List[Replacement], counts: List[int]
) -> Tuple[str, int]:
    # Initialize the total count of replacements
    total = 0
    # Initialize the updated text
    updated = text
    for idx, replacement in enumerate(replacements):
        # Apply the replacements to the text
        updated, count = replacement.pattern.subn(replacement.new_text, updated)
        # Track per-replacement counts
        if count:
            counts[idx] += count
            total += count
    # Return the updated text and the total count of replacements
    return updated, total


def _replace_in_text_chunks(
    run_texts: List[str],
    replacements: List[Replacement],
    counts: List[int],
) -> Tuple[List[str], int]:
    total = 0
    safety_limit = 10000
    iterations = 0
    texts = list(run_texts)
    
    # sourcery skip: while-guard-to-condition
    while iterations < safety_limit:
        if not any(texts):
            break
        full_text = "".join(texts)

        # Find the earliest match among all replacement patterns.
        earliest_match = None
        earliest_index = -1
        for idx, replacement in enumerate(replacements):
            match = replacement.pattern.search(full_text)
            if match and (
                earliest_match is None or match.start() < earliest_match.start()
            ):
                earliest_match = match
                earliest_index = idx

        if earliest_match is None:
            break

        start, end = earliest_match.span()
        if start == end:
            break

        run_spans: List[Tuple[int, int]] = []
        pos = 0
        for text in texts:
            run_spans.append((pos, pos + len(text)))
            pos += len(text)

        start_run = _find_run_index(run_spans, start)
        end_run = _find_run_index(run_spans, end - 1)

        start_run_start = run_spans[start_run][0]
        end_run_start = run_spans[end_run][0]

        prefix = texts[start_run][: start - start_run_start]
        suffix = texts[end_run][end - end_run_start :]

        texts[start_run] = prefix + replacements[earliest_index].new_text + suffix
        for idx in range(start_run + 1, end_run + 1):
            texts[idx] = ""

        counts[earliest_index] += 1
        total += 1
        iterations += 1

    return texts, total


def apply_xml_replacements(
    docx_path: Path,
    replacements: List[Replacement],
    ignore_case: bool,
    include_headers_footers: bool,
    apply_changes: bool,
    debug_log: Path | None = None,
    doc_label: str | None = None,
) -> List[int]:
    # Replace across XML text nodes without reserializing the full XML tree.
    # This keeps the original XML structure intact and avoids Word warnings.
    # For XML we build flexible patterns that tolerate whitespace variations.
    flags = re.IGNORECASE if ignore_case else 0
    xml_replacements: List[Replacement] = []
    for replacement in replacements:
        pattern_text = re.escape(replacement.old_text)
        # Allow any whitespace (space/tab/NBSP) where a space exists in the lookup.
        pattern_text = pattern_text.replace(r"\ ", r"[\s\u00A0]*")
        pattern = re.compile(pattern_text, flags=flags)
        xml_replacements.append(
            Replacement(
                old_text=replacement.old_text,
                new_text=replacement.new_text,
                pattern=pattern,
            )
        )
    counts = [0] * len(replacements)

    # Compile the text node pattern
    text_node_pattern = re.compile(
        r"(<(?P<tag>(?:w|a):t|w:instrText)\b[^>]*>)"
        r"(?P<text>.*?)"
        r"(</(?P=tag)>)",
        re.DOTALL,
    )
    # Compile the fallback pattern
    fallback_pattern = re.compile(
        r"<mc:Fallback[^>]*>.*?</mc:Fallback>", re.DOTALL
    )

    # Define a function to determine if a file should be processed
    def should_process(filename: str) -> bool:
        # Only operate on Word XML parts, optionally skipping headers/footers.
        if not filename.startswith("word/") or not filename.endswith(".xml"):
            return False
        return (
            include_headers_footers
            or not filename.startswith("word/header")
            and not filename.startswith("word/footer")
        )

    # Define a function to determine if a position is in a fallback range
    def is_in_fallback(position: int, ranges: List[Tuple[int, int]]) -> bool:
        for start, end in ranges:
            if start <= position < end:
                return True
        return False

    # Define a function to extract the text nodes from the XML text
    def extract_text_nodes(
        xml_text: str,
    ) -> Tuple[List[re.Match[str]], List[str], List[str]]:
        fallback_ranges = [
            (match.start(), match.end())
            for match in fallback_pattern.finditer(xml_text)
        ]
        matches: List[re.Match[str]] = []
        texts: List[str] = []
        raw_texts: List[str] = []
        for match in text_node_pattern.finditer(xml_text):
            if is_in_fallback(match.start(), fallback_ranges):
                continue
            matches.append(match)
            raw_text = match.group("text")
            decoded = unescape(raw_text).replace("\u00A0", " ")
            texts.append(decoded)
            raw_texts.append(raw_text)
        return matches, texts, raw_texts

    # Define a function to encode the XML text
    def encode_xml_text(text: str) -> str:
        encoded_parts: List[str] = []
        for ch in text:
            code = ord(ch)
            if ch == "&":
                encoded_parts.append("&amp;")
            elif ch == "<":
                encoded_parts.append("&lt;")
            elif ch == ">":
                encoded_parts.append("&gt;")
            elif code in {0x9, 0xA, 0xD} or 0x20 <= code <= 0xD7FF or 0xE000 <= code <= 0xFFFD:
                encoded_parts.append(ch)
            else:
                encoded_parts.append(f"&#x{code:X};")
        return "".join(encoded_parts)

    # Define a function to log the XML debug information
    def log_xml_debug(
        part_name: str,
        match: re.Match[str],
        node_index: int,
        old_raw: str,
        old_decoded: str,
        new_decoded: str,
        new_encoded: str,
    ) -> None:
        if not debug_log:
            return
        debug_log.parent.mkdir(parents=True, exist_ok=True)
        exists = debug_log.exists()
        with debug_log.open("a", newline="", encoding="utf-8") as handle:
            writer = csv.writer(handle)
            if not exists:
                writer.writerow(
                    [
                        "document",
                        "part",
                        "tag",
                        "node_index",
                        "old_raw",
                        "old_decoded",
                        "new_decoded",
                        "new_encoded",
                    ]
                )
            writer.writerow(
                [
                    doc_label or "",
                    part_name,
                    match.group("tag"),
                    node_index,
                    old_raw,
                    old_decoded,
                    new_decoded,
                    new_encoded,
                ]
            )

    # Define a function to rebuild the XML text
    def rebuild_xml(
        xml_text: str,
        matches: List[re.Match[str]],
        updated_decoded: List[str],
        original_decoded: List[str],
        raw_texts: List[str],
        part_name: str,
    ) -> str:
        """Splice updated node contents back into the original XML string.

        Only nodes whose decoded text actually changed are re-encoded.
        Everything else — including bytes between nodes — is taken verbatim
        from xml_text so no structural XML is ever rewritten.
        """
        parts: List[str] = []
        last_end = 0
        for node_index, match in enumerate(matches):
            # Append everything between the previous node and this one, unchanged.
            parts.append(xml_text[last_end : match.start()])

            new_decoded = updated_decoded[node_index]
            orig_decoded = original_decoded[node_index]
            raw = raw_texts[node_index]

            if new_decoded == orig_decoded:
                # Unchanged: splice the whole original element back verbatim.
                parts.append(match.group(0))
            else:
                # Changed: encode new text and rebuild element with correct groups.
                # group(1) = open tag, group(4) = close tag (group(3) = inner text).
                encoded = encode_xml_text(new_decoded)
                start_tag = match.group(1)
                end_tag = match.group(4)
                parts.append(f"{start_tag}{encoded}{end_tag}")

                log_xml_debug(
                    part_name,
                    match,
                    node_index,
                    raw,
                    orig_decoded,
                    new_decoded,
                    encoded,
                )

            last_end = match.end()

        # Append everything after the last node, unchanged.
        parts.append(xml_text[last_end:])
        return "".join(parts)

    if not apply_changes:
        # Read-only mode: count matches without writing any files.
        with zipfile.ZipFile(docx_path, "r") as source:
            for info in source.infolist():
                if not should_process(info.filename):
                    continue
                # errors="replace" so no bytes silently vanish during decode.
                xml_text = source.read(info.filename).decode("utf-8", errors="replace")
                _, run_texts, _ = extract_text_nodes(xml_text)
                local_counts = [0] * len(xml_replacements)
                _replace_in_text_chunks(run_texts, xml_replacements, local_counts)
                for idx, value in enumerate(local_counts):
                    counts[idx] += value
        return counts

    # Write a new zip to a temp file, then replace the original.
    # Use a fresh ZipInfo for every entry — stale size/CRC/extra fields from
    # the source (e.g. NTFS timestamps, ZIP64 fields) cause Word to refuse to
    # open the file when the data size has changed.
    temp_path = docx_path.with_suffix(f"{docx_path.suffix}.tmp")
    with zipfile.ZipFile(docx_path, "r") as source:
        with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as target:
            for info in source.infolist():
                raw_data = source.read(info.filename)

                if not should_process(info.filename):
                    # Non-XML or excluded part: copy verbatim with a clean ZipInfo.
                    clean_info = zipfile.ZipInfo(
                        filename=info.filename,
                        date_time=info.date_time,
                    )
                    clean_info.compress_type = info.compress_type
                    target.writestr(clean_info, raw_data)
                    continue

                xml_text = raw_data.decode("utf-8", errors="replace")
                matches, run_texts, raw_texts = extract_text_nodes(xml_text)
                local_counts = [0] * len(xml_replacements)
                updated_texts, replacements_made = _replace_in_text_chunks(
                    run_texts, xml_replacements, local_counts
                )
                for idx, value in enumerate(local_counts):
                    counts[idx] += value

                if replacements_made and matches:
                    new_xml = rebuild_xml(
                        xml_text, matches, updated_texts,
                        run_texts, raw_texts, info.filename,
                    )
                    new_data = new_xml.encode("utf-8")
                else:
                    new_data = raw_data

                # Always use a fresh ZipInfo with no stale metadata.
                clean_info = zipfile.ZipInfo(
                    filename=info.filename,
                    date_time=info.date_time,
                )
                clean_info.compress_type = zipfile.ZIP_DEFLATED
                target.writestr(clean_info, new_data)

    _replace_with_retries(temp_path, docx_path)
    return counts

# Define a function to replace the source path with the target path with retries
def _replace_with_retries(
    source_path: Path,
    target_path: Path,
    attempts: int = 12,
    delay_seconds: float = 0.5,
) -> None:
    # Windows and sync tools can lock files briefly; retry the move.
    last_error: Exception | None = None
    for _ in range(attempts):
        try:
            source_path.replace(target_path)
            return
        except PermissionError as exc:
            last_error = exc
            time.sleep(delay_seconds)
    raise PermissionError(
        f"Failed to replace output file after {attempts} attempts. "
        "Close any apps or preview panes using the file and try again. "
        f"Updated file is at: {source_path}"
    ) from last_error


# Define a function to get the working path
def _working_path(output_path: Path) -> Path:
    # Pick a stable working path to avoid OneDrive/preview locks.
    for idx in range(100):
        suffix = f".working-{idx}" if idx else ".working"
        candidate = output_path.with_name(f"{output_path.name}{suffix}")
        if not candidate.exists():
            return candidate
    unique_suffix = uuid.uuid4().hex[:8]
    return output_path.with_name(f"{output_path.name}.working-{unique_suffix}")


# Define a function to find the run index for a character position
def _find_run_index(run_spans: List[Tuple[int, int]], position: int) -> int:
    for idx, (start, end) in enumerate(run_spans):
        if start <= position < end:
            return idx
    return next(
        (
            idx
            for idx in range(len(run_spans) - 1, -1, -1)
            if run_spans[idx][0] != run_spans[idx][1]
        ),
        0,
    )


# Define a function to replace the text in the paragraph
def replace_in_paragraph(
    paragraph, replacements: List[Replacement], counts: List[int]
) -> int:
    # Initialize the total count of replacements
    total = 0
    # Iterate through the runs in the paragraph
    for run in paragraph.runs:
        # If the run text is None, continue
        if not run.text:
            continue
        # Apply the replacements to the text
        updated, count = apply_replacements(run.text, replacements, counts)
        # If the count of replacements is greater than 0, update the run text
        if count:
            run.text = updated
            # Increment the total count of replacements
            total += count
    # Return the total count of replacements
    return total


# Define a function to replace the text across runs in the paragraph
def replace_in_paragraph_join_runs(
    paragraph, replacements: List[Replacement], counts: List[int]
) -> int:
    # Initialize the total count of replacements
    total = 0
    safety_limit = 10000
    iterations = 0

    # While the iterations are less than the safety limit
    while iterations < safety_limit:
        # Get the runs in the paragraph
        runs = paragraph.runs
        # Get the text in the runs
        run_texts = [run.text or "" for run in runs]
        # If there are no runs, break
        if not any(run_texts):
            break
        # Treat the paragraph as one string so matches can cross runs
        full_text = "".join(run_texts)

        # Find the earliest match among all replacement patterns
        earliest_match = None
        earliest_index = -1
        for idx, replacement in enumerate(replacements):
            # Search for the match in the full text
            match = replacement.pattern.search(full_text)
            # If the match is found and it is the earliest match, set the earliest match and index
            if match and (
                earliest_match is None or match.start() < earliest_match.start()
            ):
                earliest_match = match
                earliest_index = idx

        # If the earliest match is not found, break
        if earliest_match is None:
            break

        # Get the start and end of the earliest match
        start, end = earliest_match.span()
        # If the start and end are the same, break
        if start == end:
            break

        # Map character positions back to run boundaries
        run_spans: List[Tuple[int, int]] = []
        pos = 0
        for text in run_texts:
            # Add the position and the length of the text to the run spans
            run_spans.append((pos, pos + len(text)))
            pos += len(text)

        # Find the start run index
        start_run = _find_run_index(run_spans, start)
        # Find the end run index
        end_run = _find_run_index(run_spans, end - 1)

        # Get the start of the start run
        start_run_start = run_spans[start_run][0]
        # Get the start of the end run
        end_run_start = run_spans[end_run][0]

        # Replace the match inside the start run and clear intervening runs
        prefix = run_texts[start_run][: start - start_run_start]
        suffix = run_texts[end_run][end - end_run_start :]

        # Update the text in the start run
        runs[start_run].text = (
            prefix + replacements[earliest_index].new_text + suffix
        )
        # Clear the intervening runs
        for idx in range(start_run + 1, end_run + 1):
            runs[idx].text = ""

        # Increment the count of the earliest index
        counts[earliest_index] += 1
        # Increment the total count of replacements
        total += 1
        # Increment the iterations
        iterations += 1

    # Return the total count of replacements
    return total


# Define a function to process the document
def process_document(
    input_path: Path,
    output_path: Path,
    replacements: List[Replacement],
    include_headers_footers: bool,
    dry_run: bool,
    join_runs: bool,
    xml_replace: bool,
    ignore_case: bool,
    xml_debug_path: Path | None,
    doc_label: str,
) -> Tuple[bool, int, List[int]]:
    # Load the document from the input path
    doc = Document(str(input_path))
    # Initialize the total count of replacements to 0
    total_replacements = 0
    # Initialize the per-replacement counts to 0
    counts = [0] * len(replacements)
    # Define a function to process the container
    def process_container(container) -> None:
        # Update the total count of replacements
        nonlocal total_replacements
        # Iterate through the paragraphs in the container
        for paragraph in iter_paragraphs(container):
            # Replace the text in the paragraph
            if join_runs:
                # Replace the text in the paragraph with the join runs
                total_replacements += replace_in_paragraph_join_runs(
                    paragraph, replacements, counts
                )
            else:
                total_replacements += replace_in_paragraph(
                    paragraph, replacements, counts
                )

    # Process the document
    process_container(doc)

    # If include headers and footers, process the headers and footers
    if include_headers_footers:
        # Initialize a set of seen elements to an empty set
        seen = set()
        # Iterate through the sections in the document
        for section in doc.sections:
            # Get the header and footer groups for the section
            # The header and footer groups are the header, first page header, even page header, footer, first page footer, and even page footer
            header_footer_groups = [
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ]
            # Iterate through the header and footer groups
            for header_footer in header_footer_groups:
                # If the header footer is None, continue
                if header_footer is None:
                    continue
                # Get the element id
                element_id = id(header_footer._element)
                # If the element id is not in the set of seen elements, add it to the set of seen elements
                if element_id not in seen:
                    seen.add(element_id)
                    # Process the container
                    process_container(header_footer)

    # Optional XML pass: required for text boxes/shapes not exposed by python-docx.
    xml_counts = [0] * len(replacements)
    # If dry run is enabled
    if dry_run:
        # If XML replace is enabled
        if xml_replace:
            # Count XML matches without writing output
            xml_counts = apply_xml_replacements(
                input_path,
                replacements,
                ignore_case=ignore_case,
                include_headers_footers=include_headers_footers,
                apply_changes=False,
                debug_log=xml_debug_path,
                doc_label=doc_label,
            )
            # Iterate through the XML counts
            for idx, count in enumerate(xml_counts):
                # If the count is greater than 0, increment the total count of replacements
                if count:
                    counts[idx] += count
                    total_replacements += count
        # Set the changed flag to True if the total count of replacements is greater than 0
        changed = total_replacements > 0
        # Return the changed flag, the total count of replacements, and the counts
        return changed, total_replacements, counts

    # Ensure output exists before applying XML replacement
    output_path.parent.mkdir(parents=True, exist_ok=True)
    work_path = output_path
    # If XML replace is enabled
    if xml_replace:
        # Write to a working file first to avoid OneDrive/preview locks
        work_path = _working_path(output_path)

    # Write the updated docx (or copy original when no text changes)
    if total_replacements > 0:
        doc.save(str(work_path))
    else:
        # If the work path is not the same as the input path, copy the input path to the work path
        if work_path.resolve() != input_path.resolve():
            shutil.copy2(input_path, work_path)

    # If XML replace is enabled
    if xml_replace:
        # Apply XML replacements to the working file and move into place
        xml_counts = apply_xml_replacements(
            work_path,
            replacements,
            ignore_case=ignore_case,
            include_headers_footers=include_headers_footers,
            apply_changes=True,
            debug_log=xml_debug_path,
            doc_label=doc_label,
        )
        # Iterate through the XML counts
        for idx, count in enumerate(xml_counts):
            # If the count is greater than 0, increment the total count of replacements
            if count:
                counts[idx] += count
                total_replacements += count
        # If the work path is not the same as the output path, replace the work path with the output path with retries
        if work_path != output_path:
            _replace_with_retries(work_path, output_path)

    # Set the changed flag to True if the total count of replacements is greater than 0
    changed = total_replacements > 0
    # Return the changed flag, the total count of replacements, and the counts
    return changed, total_replacements, counts


# Define a function to main
def main() -> int:
    # Create the parser for the command line arguments
    parser = argparse.ArgumentParser(
        description="Mass find/replace in Word .docx templates using an Excel lookup table."
    )
    # Add the excel argument to the parser
    parser.add_argument(
        "--excel",
        default=DEFAULT_EXCEL_PATH,
        help="Path to Excel lookup table (.xlsx).",
    )
    # Add the sheet argument to the parser
    parser.add_argument(
        "--sheet",
        default=DEFAULT_SHEET_NAME,
        help="Worksheet name containing Old/New values.",
    )
    # Add the input-dir argument to the parser
    parser.add_argument(
        "--input-dir",
        required=True,
        help="Directory containing Word .docx templates.",
    )
    # Add the output-dir argument to the parser
    parser.add_argument(
        "--output-dir",
        help=(
            "Directory to write updated templates. "
            "Defaults to <input-dir>_updated if not using --in-place."
        ),
    )
    # Add the in-place argument to the parser
    parser.add_argument(
        "--in-place",
        action="store_true",
        help="Overwrite templates in place (use with caution).",
    )
    # Add the old-col argument to the parser
    parser.add_argument(
        "--old-col",
        help="Old value column header or 1-based index (e.g., 'Old' or '1').",
    )
    # Add the new-col argument to the parser
    parser.add_argument(
        "--new-col",
        help="New value column header or 1-based index (e.g., 'New' or '2').",
    )
    # Add the literal argument to the parser
    parser.add_argument(
        "--literal",
        action="store_true",
        help="Treat Old values as literal text instead of regex patterns.",
    )
    # Add the ignore-case argument to the parser
    parser.add_argument(
        "--ignore-case",
        action="store_true",
        help="Case-insensitive matching.",
    )
    # Add the skip-headers-footers argument to the parser
    parser.add_argument(
        "--skip-headers-footers",
        action="store_true",
        help="Skip replacements in headers and footers.",
    )
    # Add the dry-run argument to the parser    
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Preview matches without writing files.",
    )
    # Add the xml-replace argument to the parser
    parser.add_argument(
        "--xml-replace",
        action="store_true",
        help="Also replace placeholders inside docx XML (text boxes/shapes).",
    )
    # Add the xml-debug argument to the parser
    parser.add_argument(
        "--xml-debug",
        nargs="?",
        const="__default__",
        help=(
            "Write XML debug log to path (default: xml_debug_log.csv in output dir)."
        ),
    )
    # Add the join-runs argument to the parser
    parser.add_argument(
        "--join-runs",
        action="store_true",
        help=(
            "Allow replacements spanning multiple Word runs. "
            "Uses the first run's formatting for the replaced span."
        ),
    )
    # Add the report argument to the parser
    parser.add_argument(
        "--report",
        help=(
            "CSV report path for per-document replacements. "
            "Defaults to replacement_report.csv in the output directory."
        ),
    )
    # Parse the arguments
    args = parser.parse_args()
    # Resolve the input directory

    # If XML replace is enabled and literal is not enabled, print a warning
    if args.xml_replace and not args.literal:
        print(
            "Warning: --xml-replace uses literal matching for XML text nodes; "
            "regex patterns are not applied in XML.",
            file=sys.stderr,
        )

    # Resolve the input directory
    input_dir = Path(args.input_dir).resolve()
    # If the input directory does not exist or is not a directory, raise an error
    if not input_dir.exists() or not input_dir.is_dir():
        # If the input directory does not exist or is not a directory, raise an error
        print(f"Input directory not found: {input_dir}", file=sys.stderr)
        return 1

    # Initialize the output directory to the input directory or the input directory with _updated appended
    output_dir = Path(args.output_dir).resolve() if args.output_dir else Path(f"{input_dir}_updated")
    # Set the report path to the report path or the input directory with replacement_report.csv if dry run is not enabled or the output directory
    report_path = Path(args.report).resolve() if args.report else (input_dir if args.dry_run else output_dir) / "replacement_report.csv"
    # If the report path does not exist, create it
    if not report_path.exists():
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.touch()

    # Initialize the XML debug path to None
    xml_debug_path: Path | None = None
    # If XML debug is enabled
    if args.xml_debug:
        if args.xml_debug == "__default__":
            # Set the XML debug path to the input directory with xml_debug_log.csv
            xml_debug_path = (input_dir if args.dry_run else output_dir) / "xml_debug_log.csv"
        else:
            # Set the XML debug path to the XML debug path
            xml_debug_path = Path(args.xml_debug).resolve()

    # Load the lookup table
    replacements = load_lookup_table(
        # Resolve the excel path
        Path(args.excel).resolve(),
        # Get the sheet name
        args.sheet,
        # Get the old column
        args.old_col,
        # Get the new column
        args.new_col,
        # If literal is not enabled, use the regex flag otherwise use the escape flag
        use_regex=not args.literal,
        # If ignore case is enabled, use the ignore case flag
        ignore_case=args.ignore_case,
    )
    # Load the lookup table
    replacements = load_lookup_table(
        # Resolve the excel path
        Path(args.excel).resolve(),
        # Get the sheet name
        args.sheet,
        # Get the old column
        args.old_col,
        # Get the new column
        args.new_col,
        # If literal is not enabled, use the regex flag otherwise use the escape flag
        use_regex=not args.literal,
        # If ignore case is enabled, use the ignore case flag
        ignore_case=args.ignore_case,
    )
    # Iterate through the files in the input directory
    for file_path in input_dir.rglob("*.docx"):
        # If the file name starts with ~$, continue
        if not file_path.name.startswith("~$"):
            # Process the document
            changed, count, counts = process_document(
                # Get the file path
                file_path,
                # Get the target path
                output_dir / file_path.relative_to(input_dir),
                # Get the replacements
                replacements,
                # If skip headers and footers is enabled, include headers and footers
                include_headers_footers=not args.skip_headers_footers,
                # If dry run is enabled, dry run the document
                dry_run=args.dry_run,
                # If join runs is enabled, allow run-spanning replacements
                join_runs=args.join_runs,
                # If XML replace is enabled, update XML text nodes
                xml_replace=args.xml_replace,
                # If ignore case is enabled, match case-insensitively
                ignore_case=args.ignore_case,
                # Optional XML debug log path
                xml_debug_path=xml_debug_path,
                # Label for debug logs (relative path)      
                doc_label=str(file_path.relative_to(input_dir)),
            )
            # If the document was changed, write the report to the report path
            if changed:
                # Write the report to the report path
                report_path.write_text(f"{file_path.relative_to(input_dir)},{count}\n")
            # If the document was not changed, write the report to the report path
            else:
                # Write the report to the report path
                report_path.write_text(f"{file_path.relative_to(input_dir)},0\n")   
    # Return 0 to indicate success
    return 0

# If the script is run directly, raise an error
if __name__ == "__main__":
    # Raise an error
    raise SystemExit(main())
