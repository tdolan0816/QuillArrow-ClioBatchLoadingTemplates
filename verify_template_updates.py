#!/usr/bin/env python3
"""
Verify that Old custom fields no longer exist in updated templates.

Outputs a CSV report of any Old values still found in the templates.
"""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

from docx import Document
from openpyxl import load_workbook


DEFAULT_EXCEL_PATH = (
    r"C:\Users\Tim\OneDrive - quillarrowlaw.com\Documents\ClioTemplates_CustomFields_MassUpdate"
    r"\CustomField_LookupTable.xlsx"
)
DEFAULT_SHEET_NAME = "LookupTable"


@dataclass(frozen=True)
class Replacement:
    old_text: str
    new_text: str
    pattern: re.Pattern[str]


def _normalize_header(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip().lower()


def _resolve_column_index(
    header_row: Iterable[object],
    header_map: Dict[str, int],
    column_arg: Optional[str],
    fallbacks: List[str],
    column_label: str,
) -> int:
    if column_arg:
        if column_arg.strip().isdigit():
            idx = int(column_arg.strip()) - 1
            if idx < 0 or idx >= len(header_row):
                raise ValueError(
                    f"{column_label} column index {column_arg} is out of range."
                )
            return idx
        normalized = _normalize_header(column_arg)
        if normalized in header_map:
            return header_map[normalized]
        raise ValueError(
            f"{column_label} column '{column_arg}' was not found in the header row."
        )

    for name in fallbacks:
        if name in header_map:
            return header_map[name]

    if len(header_row) >= 2:
        return 0 if column_label == "Old" else 1

    raise ValueError(
        "Lookup table must contain at least two columns for Old/New values."
    )


def load_lookup_table(
    excel_path: Path,
    sheet_name: str,
    old_col: Optional[str],
    new_col: Optional[str],
    use_regex: bool,
    ignore_case: bool,
) -> List[Replacement]:
    if not excel_path.exists():
        raise FileNotFoundError(f"Excel file not found: {excel_path}")

    workbook = load_workbook(excel_path, data_only=True, read_only=True)
    if sheet_name not in workbook.sheetnames:
        raise ValueError(
            f"Worksheet '{sheet_name}' not found. Available: {workbook.sheetnames}"
        )

    sheet = workbook[sheet_name]
    rows = sheet.iter_rows(values_only=True)
    try:
        header_row = next(rows)
    except StopIteration as exc:
        raise ValueError("Lookup table sheet is empty.") from exc

    header_map: Dict[str, int] = {}
    for idx, value in enumerate(header_row):
        normalized = _normalize_header(value)
        if normalized:
            header_map[normalized] = idx

    old_idx = _resolve_column_index(
        header_row,
        header_map,
        old_col,
        ["old", "old value", "old_value", "old values", "oldvalues"],
        "Old",
    )
    new_idx = _resolve_column_index(
        header_row,
        header_map,
        new_col,
        ["new", "new value", "new_value", "new values", "newvalues"],
        "New",
    )

    flags = re.IGNORECASE if ignore_case else 0
    replacements: List[Replacement] = []

    for row in rows:
        if row is None:
            continue
        old_value = row[old_idx] if old_idx < len(row) else None
        new_value = row[new_idx] if new_idx < len(row) else None
        if old_value is None or str(old_value).strip() == "":
            continue

        old_text = str(old_value)
        new_text = "" if new_value is None else str(new_value)
        pattern_text = old_text if use_regex else re.escape(old_text)
        try:
            pattern = re.compile(pattern_text, flags=flags)
        except re.error as exc:
            raise ValueError(
                f"Invalid regex pattern in Old value '{old_text}': {exc}"
            ) from exc

        replacements.append(
            Replacement(old_text=old_text, new_text=new_text, pattern=pattern)
        )

    if not replacements:
        raise ValueError("No replacement rows found in the lookup table.")

    return replacements


def iter_paragraphs(container) -> Iterable:
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def collect_text(container, include_headers_footers: bool) -> str:
    texts: List[str] = []
    for paragraph in iter_paragraphs(container):
        if paragraph.text:
            texts.append(paragraph.text)

    if include_headers_footers:
        seen = set()
        for section in container.sections:
            header_footer_groups = [
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ]
            for header_footer in header_footer_groups:
                element_id = id(header_footer._element)
                if element_id in seen:
                    continue
                seen.add(element_id)
                for paragraph in iter_paragraphs(header_footer):
                    if paragraph.text:
                        texts.append(paragraph.text)

    return "\n".join(texts)


def count_occurrences(text: str, pattern: re.Pattern[str]) -> int:
    if not text:
        return 0
    return len(pattern.findall(text))


def scan_docx_text(
    doc: Document, replacements: List[Replacement], include_headers_footers: bool
) -> Dict[str, int]:
    text = collect_text(doc, include_headers_footers)
    counts: Dict[str, int] = {}
    for replacement in replacements:
        counts[replacement.old_text] = count_occurrences(text, replacement.pattern)
    return counts


def scan_docx_xml(
    docx_path: Path, replacements: List[Replacement]
) -> Dict[str, int]:
    counts: Dict[str, int] = {r.old_text: 0 for r in replacements}
    with zipfile.ZipFile(docx_path) as archive:
        for name in archive.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            xml_text = archive.read(name).decode("utf-8", errors="ignore")
            plain_text = re.sub(r"<[^>]+>", "", xml_text)
            for replacement in replacements:
                counts[replacement.old_text] += count_occurrences(
                    plain_text, replacement.pattern
                )
    return counts


def write_report(report_path: Path, rows: List[List[str]]) -> None:
    report_path.parent.mkdir(parents=True, exist_ok=True)
    with report_path.open("w", encoding="utf-8", newline="") as handle:
        handle.write(
            "document,old_value,new_value,count_docx,count_xml\n"
        )
        for row in rows:
            handle.write(",".join(row) + "\n")


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Verify that Old custom fields no longer exist in updated templates."
    )
    parser.add_argument("--input-dir", required=True)
    parser.add_argument("--excel", default=DEFAULT_EXCEL_PATH)
    parser.add_argument("--sheet", default=DEFAULT_SHEET_NAME)
    parser.add_argument("--old-col")
    parser.add_argument("--new-col")
    parser.add_argument("--literal", action="store_true")
    parser.add_argument("--ignore-case", action="store_true")
    parser.add_argument("--skip-headers-footers", action="store_true")
    parser.add_argument("--deep-scan", action="store_true")
    parser.add_argument(
        "--report",
        help="CSV output report path (default: verification_report.csv in input dir).",
    )
    parser.add_argument(
        "--report-all",
        action="store_true",
        help="Include all combinations (even zero counts) in the report.",
    )
    parser.add_argument(
        "--fail-on-findings",
        action="store_true",
        help="Exit with code 2 if any Old values are found.",
    )
    args = parser.parse_args()

    input_dir = Path(args.input_dir).resolve()
    if not input_dir.exists():
        print(f"Input directory not found: {input_dir}", file=sys.stderr)
        return 1

    replacements = load_lookup_table(
        Path(args.excel).resolve(),
        args.sheet,
        args.old_col,
        args.new_col,
        use_regex=not args.literal,
        ignore_case=args.ignore_case,
    )

    report_path = (
        Path(args.report).resolve()
        if args.report
        else input_dir / "verification_report.csv"
    )

    docx_files = [
        path
        for path in input_dir.rglob("*.docx")
        if not path.name.startswith("~$")
    ]

    total_findings = 0
    error_files: List[str] = []
    report_rows: List[List[str]] = []

    for docx_path in docx_files:
        rel_path = str(docx_path.relative_to(input_dir))
        try:
            doc = Document(str(docx_path))
        except Exception as exc:
            error_files.append(f"{rel_path}: {exc}")
            continue

        counts_docx = scan_docx_text(
            doc, replacements, include_headers_footers=not args.skip_headers_footers
        )
        counts_xml = (
            scan_docx_xml(docx_path, replacements) if args.deep_scan else {}
        )

        for replacement in replacements:
            count_docx = counts_docx.get(replacement.old_text, 0)
            count_xml = counts_xml.get(replacement.old_text, 0)
            if count_docx > 0 or count_xml > 0:
                total_findings += 1
            if args.report_all or count_docx > 0 or count_xml > 0:
                report_rows.append(
                    [
                        rel_path,
                        replacement.old_text,
                        replacement.new_text,
                        str(count_docx),
                        str(count_xml) if args.deep_scan else "",
                    ]
                )

    write_report(report_path, report_rows)

    print(f"Scanned {len(docx_files)} templates.")
    print(f"Findings: {total_findings} old-value matches.")
    print(f"Report: {report_path}")
    if error_files:
        print("Errors:")
        for entry in error_files:
            print(f"- {entry}")

    if total_findings > 0 and args.fail_on_findings:
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
