#!/usr/bin/env python3
"""
Inventory custom fields in templates.

Outputs a CSV or XLSX report with:
Template_Name, Custom_Field, CF_Count
"""

from __future__ import annotations

import argparse
import csv
import html
import re
import zipfile
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

from docx import Document
from openpyxl import Workbook


def iter_paragraphs(container) -> Iterable:
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def collect_text(doc: Document, include_headers_footers: bool) -> str:
    chunks: List[str] = []
    for paragraph in iter_paragraphs(doc):
        if paragraph.text:
            chunks.append(paragraph.text)

    if include_headers_footers:
        seen = set()
        for section in doc.sections:
            header_footer_groups = [
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ]
            for header_footer in header_footer_groups:
                element_id = id(header_footer._element)
                if element_id in seen:
                    continue
                seen.add(element_id)
                for paragraph in iter_paragraphs(header_footer):
                    if paragraph.text:
                        chunks.append(paragraph.text)

    return "\n".join(chunks)


def collect_text_from_xml(docx_path: Path) -> str:
    chunks: List[str] = []
    with zipfile.ZipFile(docx_path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            if not (
                name == "word/document.xml"
                or name.startswith("word/header")
                or name.startswith("word/footer")
                or name in {"word/footnotes.xml", "word/endnotes.xml", "word/comments.xml"}
            ):
                continue
            xml_text = archive.read(name).decode("utf-8", errors="ignore")
            plain = re.sub(r"<[^>]+>", "", xml_text)
            chunks.append(html.unescape(plain))
    return "\n".join(chunks)


def count_fields(text: str, pattern: re.Pattern[str]) -> Dict[str, int]:
    counts: Dict[str, int] = {}
    for match in pattern.findall(text):
        counts[match] = counts.get(match, 0) + 1
    return counts


def write_csv(report_path: Path, rows: List[Tuple[str, str, int]]) -> None:
    report_path.parent.mkdir(parents=True, exist_ok=True)
    with report_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(["Template_Name", "Custom_Field", "CF_Count"])
        writer.writerows(rows)


def write_xlsx(report_path: Path, rows: List[Tuple[str, str, int]]) -> None:
    report_path.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Custom_Field_Inventory"
    sheet.append(["Template_Name", "Custom_Field", "CF_Count"])
    for row in rows:
        sheet.append(list(row))
    workbook.save(str(report_path))


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Inventory custom fields in templates."
    )
    parser.add_argument("--input-dir", required=True)
    parser.add_argument(
        "--pattern",
        default=r"<<.*?>>",
        help="Regex pattern to identify custom fields.",
    )
    parser.add_argument(
        "--output",
        help="Output file path (.csv or .xlsx). Default: custom_field_inventory.xlsx in input dir.",
    )
    parser.add_argument(
        "--deep-scan",
        action="store_true",
        help="Scan underlying XML to catch fields in text boxes and shapes.",
    )
    parser.add_argument(
        "--skip-headers-footers",
        action="store_true",
        help="Skip headers and footers when scanning with python-docx.",
    )
    args = parser.parse_args()

    input_dir = Path(args.input_dir).resolve()
    if not input_dir.exists():
        print(f"Input directory not found: {input_dir}")
        return 1

    output_path = (
        Path(args.output).resolve()
        if args.output
        else input_dir / "custom_field_inventory.xlsx"
    )
    pattern = re.compile(args.pattern, re.DOTALL)

    rows: List[Tuple[str, str, int]] = []
    docx_files = [
        path
        for path in input_dir.rglob("*.docx")
        if not path.name.startswith("~$")
    ]

    for docx_path in docx_files:
        template_name = docx_path.name
        try:
            if args.deep_scan:
                text = collect_text_from_xml(docx_path)
            else:
                doc = Document(str(docx_path))
                text = collect_text(
                    doc, include_headers_footers=not args.skip_headers_footers
                )
        except Exception as exc:
            print(f"Skipping {template_name}: {exc}")
            continue

        counts = count_fields(text, pattern)
        for field, count in sorted(counts.items()):
            rows.append((template_name, field, count))

    if output_path.suffix.lower() == ".csv":
        write_csv(output_path, rows)
    else:
        write_xlsx(output_path, rows)

    print(f"Scanned {len(docx_files)} templates.")
    print(f"Report written to: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
